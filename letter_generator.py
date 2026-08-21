"""Generates a personalized 400-500 word cover letter (DOCX) for a job.

Four paragraphs: intro hook -> detailed "why this role" analysis (dynamic
on job keywords, shows the posting was actually read) -> "why this
company" (dynamic on the specific company or its sector archetype) ->
call to action / closing, in professional French ("Cordialement", not
"Best regards"). Font is Calibri (a real Word font, unlike the PDF CV
which is limited to reportlab's built-in fonts).
"""
import os
import re
import unicodedata
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from config import CANDIDATE, COMPANY_ARCHETYPES, EXPORT_DIR

NAVY = RGBColor(0x1A, 0x36, 0x5D)
FONT_NAME = "Calibri"

WHY_ROLE_RULES = [
    (["product owner"],
     "Vous recherchez un Product Owner capable de definir la vision produit et de "
     "piloter sa livraison dans un contexte agile. C'est precisement ce que j'ai fait "
     "pendant 3 ans chez SNCF Gare & Connexions : j'ai defini des roadmaps produit "
     "claires, pilote le refinement des backlogs, redige des user stories precises "
     "avec criteres d'acceptation explicites, et anime les ceremonies SAFe (PI "
     "planning, sprint planning, review, retrospective). Avec ma certification PSPO I, "
     "j'ai demontre ma capacite a arbitrer les demandes metier, optimiser la valeur "
     "produit livree, et maintenir l'engagement des parties prenantes tout en "
     "respectant les contraintes techniques. Cette experience m'a donne une "
     "comprehension profonde du cycle de vie produit et des mecaniques agiles."),
    (["amoa", "programme manager", "program manager", "transformation"],
     "Votre projet de transformation requiert un profil capable d'orchestrer des "
     "changements complexes entre metiers et IT. Mon experience chez SNCF m'a forme "
     "exactement a cela : j'ai gere une reprise de donnees massive (7 000 "
     "utilisateurs), etabli une gouvernance de la donnee robuste, anime des ateliers "
     "metier/IT, et pilote l'adoption utilisateur via des campagnes de formation "
     "ciblees (plus de 50 sessions). Je comprends les enjeux reels du terrain : les "
     "processus legacy sont complexes, les utilisateurs resistent au changement, et "
     "l'IT a ses propres contraintes techniques. J'ai appris a naviguer ces tensions, "
     "a prioriser intelligemment, et a livrer de la valeur progressivement. Ma "
     "certification SAFe m'a donne les outils pour faire passer cette approche a "
     "l'echelle sur des programmes enterprise complexes."),
    (["agile coach", "scrum master"],
     "Vous cherchez un profil capable de coacher les equipes, d'optimiser les "
     "ceremonies, et de promouvoir une culture agile durable. Ma certification "
     "Leading SAFe et mon experience de management d'une equipe pluridisciplinaire de "
     "10 personnes m'ont donne cette expertise. J'ai fait progresser les bonnes "
     "pratiques SCRUM au sein de mes equipes (daily standup, planning, "
     "retrospective), aide a identifier et lever des points de blocage, et eleve "
     "progressivement le niveau de maturite agile. J'ai egalement facilite l'adoption "
     "SAFe au niveau programme (PI planning, synchronisation entre trains), ce qui "
     "m'a montre comment faire passer les pratiques agiles a l'echelle dans des "
     "organisations complexes."),
    (["aeronautique", "aéronautique", "defense", "défense", "aviation"],
     "Votre organisation opere dans le secteur exigeant de l'aeronautique/defense, un "
     "domaine que je connais intimement via ma formation ENAC et mon experience GMAO "
     "de maintenance critique. Je comprends les enjeux uniques de ce secteur : surete "
     "de fonctionnement, tracabilite, conformite reglementaire, systemes "
     "mission-critical. Ma formation d'ingenieur m'a donne un socle technique solide "
     "(C/C++, ingenierie systeme, theorie des systemes aeronautiques), que j'ai "
     "complete par une expertise en Product Management et transformation agile. Cette "
     "combinaison rare - technique, produit et agile - me positionne pour piloter vos "
     "initiatives digitales sans jamais perdre de vue les exigences critiques du "
     "secteur."),
]
DEFAULT_WHY_ROLE = (
    "Mon parcours combine pilotage agile, gestion de programme et conduite du "
    "changement, des competences qui repondent directement aux enjeux de ce poste. "
    "Chez SNCF Gare & Connexions, j'ai pilote une solution deployee aupres de 7 000 "
    "utilisateurs, en articulant vision produit, gouvernance de la donnee et "
    "accompagnement du changement - un socle solide pour reussir rapidement dans "
    "cette fonction. Mes precedentes experiences chez BNP Paribas Real Estate "
    "(pilotage de projets agiles, reporting KPI) et Air France Industries "
    "(modelisation de processus BPMN) completent ce profil par une comprehension "
    "fine des enjeux operationnels et de la performance de processus."
)

WHY_COMPANY_ARCHETYPES = {
    "conseil": (
        "{company} represente pour moi le type d'environnement ou piloter la "
        "transformation digitale de clients varies devient le coeur du metier. Apres "
        "3 ans en mode delivery operationnel chez un client unique, j'aspire a "
        "elargir mon impact en accompagnant plusieurs organisations a transformer "
        "leur chaine de valeur via le digital. Votre culture d'innovation continue, "
        "l'exposition a des enjeux varies, et la diversite des programmes proposes "
        "correspondent precisement a mes ambitions d'evolution."
    ),
    "tech_scaleup": (
        "Rejoindre {company} - un acteur en forte croissance - m'attire fortement. "
        "Apres 3 ans dans un environnement etabli aux structures hierarchiques "
        "marquees, j'aspire a un cadre plus agile, ou les decisions se prennent "
        "rapidement et ou ma contribution impacte directement la croissance de "
        "l'entreprise. Votre culture produit et votre rythme d'execution correspondent "
        "a mes valeurs, et je suis enthousiaste a l'idee de scaler votre produit tout "
        "en apportant la rigueur methodologique acquise sur des programmes a grande "
        "echelle."
    ),
    "aero_defense": (
        "{company} est un acteur reconnu des technologies critiques pour "
        "l'aeronautique et la defense - un secteur qui m'a toujours fascine, ayant ete "
        "forme a l'ENAC et specialise en avionique. Je comprends les defis uniques de "
        "ce secteur : cycles longs, conformite stricte, exigences de surete. Ma "
        "combinaison de background technique (ingenieur ENAC), d'expertise product "
        "management et de certifications agiles me positionne de maniere ideale pour "
        "accompagner vos initiatives de modernisation sans compromis sur les "
        "exigences critiques du secteur."
    ),
    "energie_utilities": (
        "{company} conduit une transformation digitale ambitieuse pour moderniser ses "
        "operations. Ce type de programme a grande echelle m'attire particulierement. "
        "Apres avoir pilote une transformation similaire chez SNCF (7 000 "
        "utilisateurs, systemes legacy), je comprends les enjeux specifiques des "
        "grands operateurs d'infrastructures critiques. Je suis convaincu que mon "
        "experience GMAO, mon management AMOA/MOE et mon expertise en conduite du "
        "changement peuvent accelerer votre initiative de modernisation."
    ),
}
DEFAULT_WHY_COMPANY = (
    "{company} correspond precisement aux valeurs de rigueur, d'impact et "
    "d'excellence operationnelle que je recherche pour la suite de mon parcours. Je "
    "suis convaincu que mon experience de pilotage produit a grande echelle et ma "
    "double certification agile (SAFe, PSPO I) constituent des atouts concrets pour "
    "contribuer rapidement a vos enjeux strategiques."
)


def _blob(job: dict) -> str:
    return f"{job.get('job_title', '')} {job.get('job_description', '')}".lower()


def _slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode()
    text = re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_")
    return text or "Company"


def _pick_why_role(blob: str) -> str:
    for keywords, text in WHY_ROLE_RULES:
        if any(kw in blob for kw in keywords):
            return text
    return DEFAULT_WHY_ROLE


def _pick_why_company(job: dict) -> str:
    company = job.get("company", "votre entreprise")
    company_lower = company.lower()
    sector_lower = (job.get("sector") or "").lower()
    blob = _blob(job)

    for archetype, companies in COMPANY_ARCHETYPES.items():
        if any(name in company_lower for name in companies):
            return WHY_COMPANY_ARCHETYPES[archetype].format(company=company)

    sector_map = {
        "conseil": "conseil",
        "tech": "tech_scaleup",
        "saas": "tech_scaleup",
        "industrie": "aero_defense",
        "energie": "energie_utilities",
        "énergie": "energie_utilities",
        "utilities": "energie_utilities",
    }
    for key, archetype in sector_map.items():
        if key in sector_lower or key in blob:
            return WHY_COMPANY_ARCHETYPES[archetype].format(company=company)

    return DEFAULT_WHY_COMPANY.format(company=company)


def build_letter_text(job: dict) -> dict:
    blob = _blob(job)
    company = job.get("company", "votre entreprise")
    job_title = job.get("job_title", "ce poste")

    intro = (
        f"Ingenieur ENAC avec 3 ans d'experience en Product Ownership et "
        f"transformation agile, j'ai pilote la conception et la livraison d'une "
        f"solution GMAO pour 3 000 gares SNCF, impactant 7 000 utilisateurs "
        f"quotidiennement. Manager d'une equipe AMOA/MOE de 10 personnes, certifie "
        f"SAFe et PSPO I, j'ai developpe une expertise en pilotage de programmes "
        f"complexes, arbitrage des parties prenantes et conduite du changement a "
        f"l'echelle entreprise. Cette experience s'appuie sur un socle technique rare "
        f"pour un profil produit : forme a l'ingenierie systeme et au developpement "
        f"(C/C++, Python), je comprends aussi bien les enjeux metier que les "
        f"contraintes techniques d'un programme. Votre offre de {job_title} chez "
        f"{company} m'interesse vivement car elle correspond precisement a mon "
        f"profil et a mes ambitions de carriere."
    )

    why_role = _pick_why_role(blob)
    why_company = _pick_why_company(job)

    closing = (
        f"Je suis enthousiaste a l'idee de discuter de la maniere dont mon experience "
        f"en product ownership, transformation agile et management AMOA/MOE peut "
        f"contribuer au succes de {company}. J'apporterai a votre equipe une "
        f"expertise rare combinant rigueur technique et pilotage produit, une "
        f"determination pour la livraison de valeur, et une capacite eprouvee a "
        f"naviguer les complexites des programmes enterprise. Ma double certification "
        f"agile (SAFe, PSPO I), mon experience multi-secteurs (ferroviaire, finance, "
        f"aeronautique) et ma pratique de l'anglais professionnel (TOEIC B2) "
        f"completent un profil immediatement operationnel.\n\n"
        f"Je reste disponible pour un entretien a votre convenance et serais ravi "
        f"d'echanger plus en detail sur vos enjeux actuels. Merci de considerer ma "
        f"candidature avec attention."
    )

    return {"intro": intro, "why_role": why_role, "why_company": why_company, "closing": closing}


def generate_letter(job: dict) -> str:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    parts = build_letter_text(job)
    company = job.get("company", "Company")

    filename = f"LM_{_slugify(company)}.docx"
    path = os.path.join(EXPORT_DIR, filename)

    doc = Document()
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = Pt(11)

    header = doc.add_paragraph()
    run = header.add_run(CANDIDATE["name"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.name = FONT_NAME
    header.add_run(
        f"\n{CANDIDATE['email']} | {CANDIDATE['phone']} | {CANDIDATE['linkedin']}"
    ).font.name = FONT_NAME

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f"Paris, le {date.today().strftime('%d/%m/%Y')}").font.name = FONT_NAME

    doc.add_paragraph()
    doc.add_paragraph(f"Objet : Candidature - {job.get('job_title', '')} chez {company}")
    doc.add_paragraph()

    doc.add_paragraph("Madame, Monsieur,")
    doc.add_paragraph()
    doc.add_paragraph(parts["intro"])
    doc.add_paragraph()
    doc.add_paragraph(parts["why_role"])
    doc.add_paragraph()
    doc.add_paragraph(parts["why_company"])
    doc.add_paragraph()
    doc.add_paragraph(parts["closing"])
    doc.add_paragraph()
    doc.add_paragraph("Cordialement,")
    sig = doc.add_paragraph()
    sig_run = sig.add_run(CANDIDATE["name"])
    sig_run.font.color.rgb = NAVY
    sig_run.font.name = FONT_NAME
    sig_run.bold = True

    doc.save(path)
    return path
